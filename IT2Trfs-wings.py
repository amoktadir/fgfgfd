import streamlit as st
import numpy as np
import pandas as pd
import graphviz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import matplotlib.pyplot as plt

# -----------------------------
# Linguistic terms (IT2TrFS)
# -----------------------------
LINGUISTIC_TERMS = {
    "strength": {
        "VLR": ((0, 0.1, 0.1, 0.1, 1, 1), (0.0, 0.1, 0.1, 0.05, 0.9, 0.9)),
        "LR": ((0.2, 0.3, 0.3, 0.4, 1, 1), (0.25, 0.3, 0.3, 0.35, 0.9, 0.9)),
        "MR": ((0.4, 0.5, 0.5, 0.6, 1, 1), (0.45, 0.5, 0.5, 0.55, 0.9, 0.9)),
        "HR": ((0.6, 0.7, 0.7, 0.8, 1, 1), (0.65, 0.7, 0.7, 0.75, 0.9, 0.9)),
        "VHR": ((0.8, 0.9, 0.9, 1, 1, 1), (0.85, 0.90, 0.90, 0.95, 0.9, 0.9)),
    },
    "influence": {
        "ELI": ((0, 0.1, 0.1, 0.2, 1, 1), (0.05, 0.1, 0.1, 0.15, 0.9, 0.9)),
        "VLI": ((0.1, 0.2, 0.2, 0.35, 1, 1), (0.15, 0.2, 0.2, 0.3, 0.9, 0.9)),
        "LI": ((0.2, 0.35, 0.35, 0.5, 1, 1), (0.25, 0.35, 0.35, 0.45, 0.9, 0.9)),
        "MI": ((0.35, 0.5, 0.5, 0.65, 1, 1), (0.40, 0.5, 0.5, 0.6, 0.9, 0.9)),
        "HI": ((0.5, 0.65, 0.65, 0.8, 1, 1), (0.55, 0.65, 0.65, 0.75, 0.9, 0.9)),
        "VHI": ((0.65, 0.80, 0.80, 0.9, 1, 1), (0.7, 0.8, 0.8, 0.85, 0.9, 0.9)),
        "EHI": ((0.8, 0.9, 0.9, 1, 1, 1), (0.85, 0.9, 0.9, 0.95, 0.9, 0.9)),
    },
}

FULL_FORMS = {
    "VLR": "Very Low Relevance",
    "LR": "Low Relevance",
    "MR": "Medium Relevance",
    "HR": "High Relevance",
    "VHR": "Very High Relevance",
    "ELI": "Extremely Low Influence",
    "VLI": "Very Low Influence",
    "LI": "Low Influence",
    "MI": "Medium Influence",
    "HI": "High Influence",
    "VHI": "Very High Influence",
    "EHI": "Extremely High Influence",
}

# -----------------------------
# UI styling
# -----------------------------
def apply_ui():
    st.set_page_config(
        page_title="IT2TrFS WINGS Analysis",
        page_icon="📊",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    st.markdown(
        """
        <style>
        .block-container { padding-top: 1.3rem; padding-bottom: 2rem; max-width: 1400px; }
        [data-testid="stSidebar"] { border-right: 1px solid rgba(49, 51, 63, .15); }
        .app-hero {
            padding: 1.2rem 1.25rem;
            border-radius: 16px;
            border: 1px solid rgba(49, 51, 63, .18);
            background: linear-gradient(135deg, rgba(99,102,241,.12), rgba(16,185,129,.10));
            margin-bottom: 1rem;
        }
        .app-hero h1 { margin: 0; line-height: 1.1; }
        .app-hero p { margin: .35rem 0 0; opacity: .86; }
        div.stButton > button { border-radius: 12px; padding: .65rem 1rem; }
        .stDataFrame, [data-testid="stDataFrame"] {
            border-radius: 14px;
            overflow: hidden;
            border: 1px solid rgba(49, 51, 63, .18);
        }
        footer { visibility: hidden; }
        </style>
        """,
        unsafe_allow_html=True,
    )

def hero():
    st.markdown(
        """
        <div class="app-hero">
          <h1>📊 IT2TrFS WINGS Method Analysis</h1>
          <p>
            Multi-expert IT2TrFS WINGS analysis for systems with interrelated components:
            total impact, receptivity, engagement, and role (cause/effect).
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

# -----------------------------
# IT2TrFS math helpers
# -----------------------------
def format_it2(it2):
    u, l = it2
    return (
        f"(({u[0]:.6f},{u[1]:.6f},{u[2]:.6f},{u[3]:.6f};{u[4]:.1f},{u[5]:.1f}), "
        f"({l[0]:.6f},{l[1]:.6f},{l[2]:.6f},{l[3]:.6f};{l[4]:.1f},{l[5]:.1f}))"
    )

def add_it2(A, B):
    Au, Al = A
    Bu, Bl = B
    new_u = (Au[0] + Bu[0], Au[1] + Bu[1], Au[2] + Bu[2], Au[3] + Bu[3], min(Au[4], Bu[4]), min(Au[5], Bu[5]))
    new_l = (Al[0] + Bl[0], Al[1] + Bl[1], Al[2] + Bl[2], Al[3] + Bl[3], min(Al[4], Bl[4]), min(Al[5], Bl[5]))
    return (new_u, new_l)

def sub_it2(A, B):
    Au, Al = A
    Bu, Bl = B
    new_u = (Au[0] - Bu[0], Au[1] - Bu[1], Au[2] - Bu[2], Au[3] - Bu[3], min(Au[4], Bu[4]), min(Au[5], Bu[5]))
    new_l = (Al[0] - Bl[0], Al[1] - Bl[1], Al[2] - Bl[2], Al[3] - Bl[3], min(Al[4], Bl[4]), min(Al[5], Bl[5]))
    return (new_u, new_l)

def mul_it2(A, B):
    Au, Al = A
    Bu, Bl = B
    new_u = (Au[0] * Bu[0], Au[1] * Bu[1], Au[2] * Bu[2], Au[3] * Bu[3], min(Au[4], Bu[4]), min(Au[5], Bu[5]))
    new_l = (Al[0] * Bl[0], Al[1] * Bl[1], Al[2] * Bl[2], Al[3] * Bl[3], min(Al[4], Bl[4]), min(Al[5], Bl[5]))
    return (new_u, new_l)

def scalar_mul_it2(k, A):
    Au, Al = A
    new_u = (k * Au[0], k * Au[1], k * Au[2], k * Au[3], Au[4], Au[5])
    new_l = (k * Al[0], k * Al[1], k * Al[2], k * Al[3], Al[4], Al[5])
    return (new_u, new_l)

def zero_it2():
    return ((0, 0, 0, 0, 1, 1), (0, 0, 0, 0, 0.9, 0.9))

def defuzz_it2(A):
    Au, Al = A
    return (Au[0] + Au[1] + Au[2] + Au[3] + Al[0] + Al[1] + Al[2] + Al[3]) / 8

def identity_it2(n):
    I_mat = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for i in range(n):
        I_mat[i][i] = ((1, 1, 1, 1, 1, 1), (1, 1, 1, 1, 1, 1))
    return I_mat

def compute_total_relation_matrix(normalized_matrix):
    n = len(normalized_matrix)
    T = [[zero_it2() for _ in range(n)] for _ in range(n)]

    Z_4d = np.zeros((2, 2, n, n, 4))
    for i in range(n):
        for j in range(n):
            Au, Al = normalized_matrix[i][j]
            Z_4d[0, 0, i, j, :] = Au[:4]
            Z_4d[0, 1, i, j, :2] = Au[4:]
            Z_4d[1, 0, i, j, :] = Al[:4]
            Z_4d[1, 1, i, j, :2] = Al[4:]

    for i in range(2):      # UMF/LMF
        for j in range(2):  # Params/Heights
            if j == 0:
                for k in range(4):
                    Z_component = Z_4d[i, j, :, :, k]
                    try:
                        T_component = Z_component @ np.linalg.pinv(np.eye(n) - Z_component)
                    except np.linalg.LinAlgError:
                        T_component = np.zeros((n, n))
                    Z_4d[i, j, :, :, k] = T_component
            else:
                pass

    for i in range(n):
        for j in range(n):
            T[i][j] = (
                (Z_4d[0, 0, i, j, 0], Z_4d[0, 0, i, j, 1], Z_4d[0, 0, i, j, 2], Z_4d[0, 0, i, j, 3], Z_4d[0, 1, i, j, 0], Z_4d[0, 1, i, j, 1]),
                (Z_4d[1, 0, i, j, 0], Z_4d[1, 0, i, j, 1], Z_4d[1, 0, i, j, 2], Z_4d[1, 0, i, j, 3], Z_4d[1, 1, i, j, 0], Z_4d[1, 1, i, j, 1]),
            )
    return T

def calculate_TI_TR(T):
    n = len(T)
    TI = [zero_it2() for _ in range(n)]
    TR = [zero_it2() for _ in range(n)]
    for i in range(n):
        for j in range(n):
            TI[i] = add_it2(TI[i], T[i][j])
            TR[j] = add_it2(TR[j], T[i][j])
    return TI, TR

def wings_method_experts(strengths_list, influence_matrices_list, weights=None):
    n = len(strengths_list[0])
    num_experts = len(strengths_list)

    if weights is None:
        weights = [1.0 / num_experts] * num_experts

    avg_sidrm = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for exp in range(num_experts):
        w = weights[exp]
        for i in range(n):
            str_w = scalar_mul_it2(w, strengths_list[exp][i])
            avg_sidrm[i][i] = add_it2(avg_sidrm[i][i], str_w)
            for j in range(n):
                if i != j:
                    inf_w = scalar_mul_it2(w, influence_matrices_list[exp][i][j])
                    avg_sidrm[i][j] = add_it2(avg_sidrm[i][j], inf_w)

    s1U = s2U = s3U = s4U = 0.0
    s1L = s2L = s3L = s4L = 0.0
    for i in range(n):
        for j in range(n):
            Au, Al = avg_sidrm[i][j]
            s1U += Au[0]; s2U += Au[1]; s3U += Au[2]; s4U += Au[3]
            s1L += Al[0]; s2L += Al[1]; s3L += Al[2]; s4L += Al[3]
    s = s1U + s2U + s3U + s4U + s1L + s2L + s3L + s4L

    Z_mat = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for i in range(n):
        for j in range(n):
            Au, Al = avg_sidrm[i][j]
            if s != 0:
                new_u = (Au[0]/s, Au[1]/s, Au[2]/s, Au[3]/s, Au[4], Au[5])
                new_l = (Al[0]/s, Al[1]/s, Al[2]/s, Al[3]/s, Al[4], Al[5])
            else:
                new_u = (0, 0, 0, 0, Au[4], Au[5])
                new_l = (0, 0, 0, 0, Al[4], Al[5])
            Z_mat[i][j] = (new_u, new_l)

    T_mat = compute_total_relation_matrix(Z_mat)
    TI, TR = calculate_TI_TR(T_mat)

    engagement = [zero_it2() for _ in range(n)]
    role = [zero_it2() for _ in range(n)]
    for i in range(n):
        engagement[i] = add_it2(TI[i], TR[i])
        role[i] = sub_it2(TI[i], TR[i])

    TI_defuzz = np.array([defuzz_it2(TI[i]) for i in range(n)])
    TR_defuzz = np.array([defuzz_it2(TR[i]) for i in range(n)])
    engagement_defuzz = np.array([defuzz_it2(engagement[i]) for i in range(n)])
    role_defuzz = np.array([defuzz_it2(role[i]) for i in range(n)])

    return {
        "average_sidrm": avg_sidrm,
        "scaling_factor": s,
        "normalized_matrix": Z_mat,
        "total_matrix": T_mat,
        "total_impact": TI,
        "total_receptivity": TR,
        "engagement": engagement,
        "role": role,
        "total_impact_defuzz": TI_defuzz,
        "total_receptivity_defuzz": TR_defuzz,
        "engagement_defuzz": engagement_defuzz,
        "role_defuzz": role_defuzz,
    }

def format_it2_df(mat, index, columns):
    df = pd.DataFrame(index=index, columns=columns)
    for i in range(len(index)):
        for j in range(len(columns)):
            df.iloc[i, j] = format_it2(mat[i][j])
    return df

# -----------------------------
# Graphviz flowchart
# -----------------------------
def generate_flowchart_for_expert(expert_data, component_names, expert_idx=None):
    graph = graphviz.Digraph(
        comment=f"IT2TrFS WINGS Flowchart - Expert {expert_idx+1}" if expert_idx is not None else "IT2TrFS WINGS Flowchart"
    )
    graph.attr(rankdir="TD", size="10,10")

    for comp_idx, comp_name in enumerate(component_names):
        strength = expert_data["strengths_linguistic"][comp_idx]
        label = f"{comp_name}\\n({strength})"
        graph.node(comp_name, label=label, shape="box", style="rounded,filled", fillcolor="lightblue", fontsize="12")

    for from_idx, from_comp in enumerate(component_names):
        for to_idx, to_comp in enumerate(component_names):
            if from_idx == to_idx:
                continue
            influence = expert_data["influence_matrix_linguistic"][from_idx][to_idx]
            if influence != "ELI":
                graph.edge(from_comp, to_comp, label=influence)

    return graph

# -----------------------------
# Word report helpers
# -----------------------------
def add_dataframe_to_doc(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns) + 1)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = ""
    for j, col in enumerate(df.columns):
        hdr[j + 1].text = str(col)

    for i, idx in enumerate(df.index):
        row = table.add_row().cells
        row[0].text = str(idx)
        for j, col in enumerate(df.columns):
            row[j + 1].text = str(df.iloc[i, j])

    doc.add_paragraph()

def create_word_report(results, component_names, n_experts=1, expert_weights=None):
    doc = Document()
    title = doc.add_heading("IT2TrFS WINGS Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Number of experts: {n_experts}")
    if expert_weights and n_experts > 1:
        doc.add_paragraph("Expert weights: " + ", ".join([f"Expert {i+1}: {w:.2f}" for i, w in enumerate(expert_weights)]))

    comp_para = doc.add_paragraph("Components analyzed: ")
    for i, name in enumerate(component_names):
        comp_para.add_run(f"{i+1}. {name}  ")

    doc.add_heading("Impact, Receptivity, Engagement, and Role Results", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Total Impact (TI)"
    hdr[2].text = "Total Receptivity (TR)"
    hdr[3].text = "Engagement (TI+TR)"
    hdr[4].text = "Role (TI-TR)"

    for i, name in enumerate(component_names):
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f"{results['total_impact_defuzz'][i]:.6f}"
        row[2].text = f"{results['total_receptivity_defuzz'][i]:.6f}"
        row[3].text = f"{results['engagement_defuzz'][i]:.6f}"
        row[4].text = f"{results['role_defuzz'][i]:.6f}"

    doc.add_heading("Component Classification", level=1)
    class_table = doc.add_table(rows=1, cols=3)
    class_table.style = "Table Grid"
    hdr = class_table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Type"
    hdr[2].text = "Role (TI-TR)"

    for i, name in enumerate(component_names):
        status = "Cause" if results["role_defuzz"][i] > 0 else "Effect"
        row = class_table.add_row().cells
        row[0].text = name
        row[1].text = status
        row[2].text = f"{results['role_defuzz'][i]:.6f}"

    doc.add_heading("Matrices", level=1)

    doc.add_heading("Average SIDRM", level=2)
    add_dataframe_to_doc(doc, format_it2_df(results["average_sidrm"], component_names, component_names))

    doc.add_heading("Normalized Matrix Z", level=2)
    add_dataframe_to_doc(doc, format_it2_df(results["normalized_matrix"], component_names, component_names))

    doc.add_heading("Total Matrix T (IT2TrFS)", level=2)
    add_dataframe_to_doc(doc, format_it2_df(results["total_matrix"], component_names, component_names))

    doc.add_heading("Interpretation of Results", level=1)
    doc.add_paragraph("Total Impact (TI) represents the outgoing influence of a component.")
    doc.add_paragraph("Total Receptivity (TR) represents the incoming influence on a component.")
    doc.add_paragraph("Engagement (TI+TR) indicates overall involvement of a component in the system.")
    doc.add_paragraph("Role (TI-TR): positive values indicate a Cause; negative values indicate an Effect.")

    return doc

def get_word_bytes(doc):
    buff = io.BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff.read()

# -----------------------------
# Better influence input (data editor)
# -----------------------------
def influence_editor(component_names, key, defaults=None):
    n = len(component_names)
    options = list(LINGUISTIC_TERMS["influence"].keys())

    if defaults is None:
        defaults = [["ELI" for _ in range(n)] for _ in range(n)]

    df = pd.DataFrame(defaults, index=component_names, columns=component_names)
    for i in range(n):
        df.iat[i, i] = "—"

    edited = st.data_editor(
        df,
        use_container_width=True,
        key=key,
        column_config={
            col: st.column_config.SelectboxColumn(
                col,
                help="Influence level (row → column)",
                options=["—"] + options,
                required=True,
            )
            for col in df.columns
        },
    )

    mat = edited.values.tolist()
    for i in range(n):
        mat[i][i] = "ELI"
    return mat

# -----------------------------
# App
# -----------------------------
def main():
    apply_ui()
    hero()

    tab_howto, tab_analysis = st.tabs(["📘 How to Use", "📊 Analysis"])

    with tab_howto:
        st.header("How to Use")
        st.markdown(
            """
            ### Overview
            This app implements the **IT2TrFS WINGS** method for analyzing complex systems with interrelated components,
            incorporating **multiple experts** and optional **expert weighting**.

            ### Steps
            1. Configure number of components and experts in the sidebar.
            2. Provide **strength/relevance** for each component.
            3. Provide **influences (row → column)** using the influence matrix editor.
            4. Run the analysis to compute **TI**, **TR**, **Engagement**, and **Role**.
            5. Export a Word report.

            ### Interpreting Role (TI - TR)
            - **Positive** → Cause
            - **Negative** → Effect
            """
        )

        with st.expander("Linguistic Terms Reference", expanded=True):
            col1, col2 = st.columns(2)

            with col1:
                st.subheader("Strength/Relevance")
                strength_df = pd.DataFrame(
                    [
                        {"Abbreviation": abbr, "Full Form": FULL_FORMS[abbr], "IT2TrFS Interval": format_it2(it2)}
                        for abbr, it2 in LINGUISTIC_TERMS["strength"].items()
                    ]
                )
                st.dataframe(strength_df, hide_index=True, use_container_width=True)

            with col2:
                st.subheader("Influence")
                influence_df = pd.DataFrame(
                    [
                        {"Abbreviation": abbr, "Full Form": FULL_FORMS[abbr], "IT2TrFS Interval": format_it2(it2)}
                        for abbr, it2 in LINGUISTIC_TERMS["influence"].items()
                    ]
                )
                st.dataframe(influence_df, hide_index=True, use_container_width=True)

    with tab_analysis:
        with st.sidebar:
            st.header("⚙️ Configuration")
            n_components = st.number_input(
                "Number of Components",
                min_value=2,
                max_value=25,
                value=3,
                help="How many components are in your system?",
            )
            n_experts = st.number_input(
                "Number of Experts",
                min_value=1,
                max_value=15,
                value=1,
                help="How many experts will provide assessments?",
            )

            st.divider()
            st.subheader("Component Names")
            component_names = []
            for i in range(n_components):
                component_names.append(
                    st.text_input(f"Component {i+1}", value=f"C{i+1}", key=f"comp_name_{i}")
                )

            expert_weights = None
            if n_experts > 1:
                st.divider()
                st.subheader("Expert Weights")
                st.caption("Weights must sum to 1.0")

                weights = []
                total_weight = 0.0
                for i in range(n_experts):
                    max_val = min(1.0, 1.0 - total_weight + (1.0 / n_experts))
                    w = st.number_input(
                        f"Expert {i+1} weight",
                        min_value=0.0,
                        max_value=max_val,
                        value=1.0 / n_experts,
                        step=0.01,
                        format="%.2f",
                        key=f"weight_{i}",
                        help=f"Maximum allowed right now: {max_val:.2f}",
                    )
                    weights.append(w)
                    total_weight += w

                st.write(f"**Current total:** {total_weight:.2f}/1.0")
                if abs(total_weight - 1.0) > 0.001:
                    st.error("Weights must sum to 1.0. Adjust weights to continue.")
                    st.stop()
                expert_weights = weights

            st.divider()
            st.info("💡 Tip: Use the terms reference in the How-To tab if you forget abbreviations.")

        with st.expander("View Linguistic Terms Mapping", expanded=False):
            col1, col2 = st.columns(2)
            with col1:
                st.subheader("Strength/Relevance")
                strength_df = pd.DataFrame(
                    [
                        {"Abbreviation": abbr, "Full Form": FULL_FORMS[abbr], "IT2TrFS Interval": format_it2(it2)}
                        for abbr, it2 in LINGUISTIC_TERMS["strength"].items()
                    ]
                )
                st.dataframe(strength_df, hide_index=True, use_container_width=True)
            with col2:
                st.subheader("Influence")
                influence_df = pd.DataFrame(
                    [
                        {"Abbreviation": abbr, "Full Form": FULL_FORMS[abbr], "IT2TrFS Interval": format_it2(it2)}
                        for abbr, it2 in LINGUISTIC_TERMS["influence"].items()
                    ]
                )
                st.dataframe(influence_df, hide_index=True, use_container_width=True)

        # Session state initialization / resizing
        if "experts_data" not in st.session_state:
            st.session_state.experts_data = {}

        for expert_idx in range(int(n_experts)):
            if expert_idx not in st.session_state.experts_data:
                st.session_state.experts_data[expert_idx] = {
                    "strengths_linguistic": ["HR" for _ in range(int(n_components))],
                    "influence_matrix_linguistic": [["ELI" for _ in range(int(n_components))] for _ in range(int(n_components))],
                }
            else:
                if len(st.session_state.experts_data[expert_idx]["strengths_linguistic"]) != int(n_components):
                    st.session_state.experts_data[expert_idx]["strengths_linguistic"] = ["HR" for _ in range(int(n_components))]
                mat = st.session_state.experts_data[expert_idx]["influence_matrix_linguistic"]
                if len(mat) != int(n_components) or any(len(row) != int(n_components) for row in mat):
                    st.session_state.experts_data[expert_idx]["influence_matrix_linguistic"] = [
                        ["ELI" for _ in range(int(n_components))] for _ in range(int(n_components))
                    ]

        st.header("👨‍💼 Expert Inputs" if n_experts > 1 else "👨‍💼 Data Input")

        # Put all inputs in a form so the page isn't constantly rerunning while editing
        with st.form("inputs_form", clear_on_submit=False):
            expert_tabs = st.tabs([f"Expert {i+1}" for i in range(int(n_experts))]) if n_experts > 1 else [st.container()]
            strengths_list = []
            influence_matrices_list = []

            for expert_idx in range(int(n_experts)):
                tab = expert_tabs[expert_idx] if n_experts > 1 else expert_tabs[0]
                with tab:
                    if n_experts > 1:
                        st.subheader(f"Expert {expert_idx+1}")
                        if expert_weights:
                            st.caption(f"Weight: {expert_weights[expert_idx]:.2f}")

                    st.subheader("Component Strengths / Relevance")
                    cols = st.columns(min(int(n_components), 6))
                    # If more than 6 components, we’ll stack in multiple rows
                    strengths_terms = st.session_state.experts_data[expert_idx]["strengths_linguistic"]

                    strengths_out = [None] * int(n_components)
                    for i in range(int(n_components)):
                        col = cols[i % len(cols)]
                        with col:
                            current = strengths_terms[i]
                            term = st.selectbox(
                                f"{component_names[i]}",
                                options=list(LINGUISTIC_TERMS["strength"].keys()),
                                index=list(LINGUISTIC_TERMS["strength"].keys()).index(current),
                                key=f"strength_{expert_idx}_{i}",
                                help=FULL_FORMS[current],
                            )
                            st.caption(FULL_FORMS[term])
                            strengths_terms[i] = term
                            strengths_out[i] = LINGUISTIC_TERMS["strength"][term]

                    st.session_state.experts_data[expert_idx]["strengths_linguistic"] = strengths_terms

                    st.divider()
                    st.subheader("Influence Matrix")
                    st.caption("Edit influences as **row → column**. Diagonal is ignored (self-strength).")

                    current_mat = st.session_state.experts_data[expert_idx]["influence_matrix_linguistic"]
                    edited_linguistic = influence_editor(
                        component_names,
                        key=f"influence_editor_{expert_idx}",
                        defaults=current_mat,
                    )
                    st.session_state.experts_data[expert_idx]["influence_matrix_linguistic"] = edited_linguistic

                    # Convert to IT2 intervals matrix for algorithm
                    inf_it2 = [[None] * int(n_components) for _ in range(int(n_components))]
                    for i in range(int(n_components)):
                        for j in range(int(n_components)):
                            if i == j:
                                inf_it2[i][j] = LINGUISTIC_TERMS["influence"]["ELI"]
                            else:
                                inf_it2[i][j] = LINGUISTIC_TERMS["influence"][edited_linguistic[i][j]]

                    strengths_list.append(strengths_out)
                    influence_matrices_list.append(inf_it2)

            submitted = st.form_submit_button("🚀 Run IT2TrFS WINGS Analysis", use_container_width=True, type="primary")

        if submitted:
            with st.spinner("Running IT2TrFS WINGS computations..."):
                results = wings_method_experts(strengths_list, influence_matrices_list, expert_weights)

            st.success("Analysis Complete!")

            # Quick KPIs
            roles = results["role_defuzz"]
            eng = results["engagement_defuzz"]
            k1, k2, k3, k4 = st.columns(4)
            k1.metric("Components", len(component_names))
            k2.metric("Experts", int(n_experts))
            k3.metric("Cause components", int((roles > 0).sum()))
            k4.metric("Top engagement", component_names[int(np.argmax(eng))])

            tab1, tab3, tab4, tab5, tab6, tab7 = st.tabs(
                ["🔗 Flowchart", "🧮 IT2TrFS Matrices", "📊 Results", "🏷️ Classification", "📈 Visualization", "📤 Export"]
            )

            with tab1:
                st.subheader("Component Interaction Flowchart")
                if n_experts > 1:
                    for expert_idx in range(int(n_experts)):
                        st.markdown(f"#### Expert {expert_idx+1}")
                        flow = generate_flowchart_for_expert(st.session_state.experts_data[expert_idx], component_names, expert_idx)
                        st.graphviz_chart(flow, use_container_width=True)
                else:
                    flow = generate_flowchart_for_expert(st.session_state.experts_data[0], component_names)
                    st.graphviz_chart(flow, use_container_width=True)

                st.info(
                    "Nodes show component strength (abbreviation). Edges show influence levels (abbreviation)."
                )

            with tab3:
                st.subheader("Average SIDRM")
                st.dataframe(format_it2_df(results["average_sidrm"], component_names, component_names), use_container_width=True)

                st.subheader("Normalized Matrix Z")
                st.dataframe(format_it2_df(results["normalized_matrix"], component_names, component_names), use_container_width=True)

                st.subheader("Total Matrix T (IT2TrFS)")
                st.dataframe(format_it2_df(results["total_matrix"], component_names, component_names), use_container_width=True)

            with tab4:
                st.subheader("Impact / Receptivity / Engagement / Role")

                df = pd.DataFrame(
                    {
                        "Component": component_names,
                        "TI (defuzz)": results["total_impact_defuzz"],
                        "TR (defuzz)": results["total_receptivity_defuzz"],
                        "Engagement (defuzz)": results["engagement_defuzz"],
                        "Role (defuzz)": results["role_defuzz"],
                        "Type": ["Cause" if r > 0 else "Effect" for r in results["role_defuzz"]],
                    }
                ).sort_values("Engagement (defuzz)", ascending=False)

                st.dataframe(df, use_container_width=True, hide_index=True)

            with tab5:
                st.subheader("Component Classification (Cause / Effect)")

                classification_df = pd.DataFrame(
                    {
                        "Component": component_names,
                        "Type": ["Cause" if r > 0 else "Effect" for r in results["role_defuzz"]],
                        "Role (defuzz)": results["role_defuzz"],
                        "Engagement (defuzz)": results["engagement_defuzz"],
                    }
                )

                cols = st.columns(3)
                for i, row in classification_df.iterrows():
                    with cols[i % 3]:
                        emoji = "➡️" if row["Type"] == "Cause" else "⬅️"
                        st.metric(
                            label=f"{emoji} {row['Component']}",
                            value=row["Type"],
                            delta=f"Role: {row['Role (defuzz)']:.6f}",
                        )

                fig, ax = plt.subplots(figsize=(10, 5))
                ax.bar(
                    classification_df["Component"],
                    classification_df["Role (defuzz)"],
                )
                ax.axhline(y=0, linestyle="--", alpha=0.6)
                ax.set_title("Component Role Values (Defuzzified)")
                ax.set_ylabel("Role (TI - TR)")
                plt.xticks(rotation=45)
                st.pyplot(fig)

            with tab6:
                st.subheader("Engagement vs Role (Cause–Effect Map)")

                fig, ax = plt.subplots(figsize=(10, 7))
                avg_eng = float(np.mean(results["engagement_defuzz"]))

                ax.axhline(y=0, linestyle="--", alpha=0.6)
                ax.axvline(x=avg_eng, linestyle="--", alpha=0.6)

                for i, name in enumerate(component_names):
                    x = float(results["engagement_defuzz"][i])
                    y = float(results["role_defuzz"][i])
                    ax.scatter(x, y, s=140, alpha=0.8)
                    ax.annotate(name, (x, y), xytext=(5, 5), textcoords="offset points")

                ax.set_xlabel("Engagement (TI+TR) (defuzz)")
                ax.set_ylabel("Role (TI-TR) (defuzz)")
                ax.set_title("Cause–Effect Diagram")
                ax.grid(True, alpha=0.25)
                st.pyplot(fig)

                st.subheader("Component Engagement Ranking")
                fig2, ax2 = plt.subplots(figsize=(10, 6))
                order = np.argsort(results["engagement_defuzz"])[::-1]
                ax2.barh(
                    [component_names[i] for i in order][::-1],
                    [results["engagement_defuzz"][i] for i in order][::-1],
                    alpha=0.85,
                )
                ax2.set_xlabel("Engagement (defuzz)")
                ax2.set_title("Ranking by Engagement")
                ax2.grid(True, alpha=0.25, axis="x")
                st.pyplot(fig2)

            with tab7:
                st.subheader("Export Word Report")
                st.caption("Download a comprehensive report including results tables and matrices.")

                doc = create_word_report(results, component_names, int(n_experts), expert_weights)
                doc_bytes = get_word_bytes(doc)

                st.download_button(
                    label="⬇️ Download Word Report (.docx)",
                    data=doc_bytes,
                    file_name="it2trfs_wings_analysis_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

if __name__ == "__main__":
    main()
